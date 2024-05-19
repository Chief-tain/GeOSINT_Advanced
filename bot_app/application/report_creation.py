import os
from datetime import datetime
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Mm

import docx
from docx.enum.dml import MSO_THEME_COLOR_INDEX

def build_report(filtered_reports_dict, begin, end, uniq, username):

        begin_report = datetime.utcfromtimestamp(begin + 86400).strftime('%Y-%m-%d')
        end_report = datetime.utcfromtimestamp(end + 86400).strftime('%Y-%m-%d')
        document = Document()

        style = document.styles['Normal']
        style.font.name = 'Times New Roman'
        style.font.size = Pt(12)
        section = document.sections[0]
        section.right_margin = Mm(10)
        section.top_margin = Mm(15)
        section.bottom_margin = Mm(10)
        section.header_distance = Mm(10)
        section.footer_distance = Mm(10)
        head = document.add_heading('Отчет по оперативной обстановке\n {} - {}\n Уникальность информации - {}%'.format(str(begin_report), str(end_report), str(uniq)))
        head.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        for key in filtered_reports_dict:
            if len(filtered_reports_dict[key]) != 0:

                par = document.add_paragraph().add_run(key.capitalize())
                par.font.size = Pt(14)
                par.bold = True

                for point in range(len(filtered_reports_dict[key])):
                    par0 = document.add_paragraph(str(point+1) + ') ')
                    par0.add_run(filtered_reports_dict[key][point].strip())
                    par0.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                    fmt0 = par0.paragraph_format
                    fmt0.first_line_indent = Mm(15)

        if not os.path.isdir(f'output\{username}'):
            os.mkdir(f'output\{username}')

        report_name = f'output\{username}\\report_{username}.docx'

        document.save(report_name)

def build_tag_report(answer, begin, end, tag, username):

    begin_report = datetime.utcfromtimestamp(begin + 86400).strftime('%Y-%m-%d')
    end_report = datetime.utcfromtimestamp(end + 86400).strftime('%Y-%m-%d')
    document = Document()

    style = document.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    section = document.sections[0]
    section.right_margin = Mm(10)
    section.top_margin = Mm(15)
    section.bottom_margin = Mm(10)
    section.header_distance = Mm(10)
    section.footer_distance = Mm(10)
    head = document.add_heading('Перечень информационно-новостных сообщений\n {} - {}\n Тематика - {}'.format(str(begin_report), str(end_report), str(tag)))
    head.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    for key in answer:
        par = document.add_paragraph().add_run(key.capitalize())
        par.font.size = Pt(14)
        par.bold = True

        for point in range(len(answer[key])):
            par0 = document.add_paragraph(str(point+1) + ') ')
            par0.add_run(answer[key][point][0].strip() + ' - ' + answer[key][point][1].strip() + ' - ' + datetime.utcfromtimestamp(answer[key][point][2] + 86400).strftime('%Y-%m-%d %H:%M'))
            par0.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            fmt0 = par0.paragraph_format
            fmt0.first_line_indent = Mm(15)

    if not os.path.isdir(f'output\{username}'):
        os.mkdir(f'output\{username}')

    tag_report_name = f'output\{username}\\tag_list_{username}.docx'

    document.save(tag_report_name)